#!/usr/bin/env python
"""Export a bilingual English-Chinese review manuscript to Word.

The output is an internal review document: manuscript text is arranged in a
two-column English/Chinese layout, followed by the current figures and tables.
"""

from __future__ import annotations

import argparse
import re
import sys
from itertools import zip_longest
from pathlib import Path
from typing import NamedTuple

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SCRIPT_PATH = Path(__file__).resolve()
PROJECT_ROOT = SCRIPT_PATH.parents[1]
SCRIPTS_DIR = SCRIPT_PATH.parent
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

import export_word_manuscript as zh_export  # noqa: E402


class MarkdownBlock(NamedTuple):
    kind: str
    text: str


class BilingualFigure(NamedTuple):
    label_en: str
    label_zh: str
    caption_en: str
    caption_zh: str
    filename: str


BILINGUAL_FIGURES = [
    BilingualFigure(
        "Figure 1",
        "图 1",
        "Public multi-cohort framework for early LUAD epithelial-myeloid niche mining.",
        "早期 LUAD 上皮-髓系 niche 挖掘的公开多队列框架。",
        "nature_redesign/nature_figure1_workflow_dataset_composition.png",
    ),
    BilingualFigure(
        "Figure 2",
        "图 2",
        "Integrated evidence prioritizes MIF-CD74/CXCR4 over broader macrophage-state axes.",
        "综合证据优先支持 MIF-CD74/CXCR4，而不是更宽泛的巨噬细胞状态轴。",
        "nature_redesign/nature_figure2_axis_evidence_perturbation.png",
    ),
    BilingualFigure(
        "Figure 3",
        "图 3",
        "Specificity auditing reframes the original SPP1 macrophage niche hypothesis.",
        "特异性审计重塑原始 SPP1 巨噬细胞 niche 假设。",
        "nature_redesign/nature_figure3_specificity_audit.png",
    ),
    BilingualFigure(
        "Figure 4",
        "图 4",
        "Patient-aware spatial progression supports source-side MIF enrichment.",
        "患者感知空间进展支持 source-side MIF 富集。",
        "nature_redesign/nature_figure4_spatial_axis_progression.png",
    ),
    BilingualFigure(
        "Figure 5",
        "图 5",
        "Score-level in-silico target prioritization ranks receptor-side CD74.",
        "score-level in-silico target prioritization 显示受体侧 CD74 排名最高。",
        "nature_redesign/nature_figure5_virtual_perturbation_priority.png",
    ),
    BilingualFigure(
        "Supplementary Figure 1",
        "补充图 1",
        "Expression-matched and tissue-density controls for source-side MIF.",
        "source-side MIF 的表达量匹配和组织密度对照。",
        "supplementary_figure_mif_spatial_controls.png",
    ),
    BilingualFigure(
        "Supplementary Figure 2",
        "补充图 2",
        "Focused orthogonal expression support in snRNA and bulk progression cohorts.",
        "snRNA 和 bulk 进展队列中的重点正交表达支持。",
        "supplementary_figure_focused_orthogonal_validation.png",
    ),
    BilingualFigure(
        "Supplementary Figure 3",
        "补充图 3",
        "Specificity refinement of the original SPP1 macrophage signature.",
        "原始 SPP1 巨噬细胞 signature 的特异性修正。",
        "supplementary_figure_spp1_signature_refinement.png",
    ),
    BilingualFigure(
        "Supplementary Figure 4",
        "补充图 4",
        "GRN-level virtual perturbation prioritization in GSE308103.",
        "GSE308103 GRN-level virtual perturbation prioritization。",
        "supplementary_figure_grn_virtual_perturbation.png",
    ),
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--english-md",
        type=Path,
        default=PROJECT_ROOT / "docs" / "manuscript_communications_biology_draft.md",
    )
    parser.add_argument(
        "--chinese-md",
        type=Path,
        default=PROJECT_ROOT / "docs" / "manuscript_communications_biology_draft_zh.md",
    )
    parser.add_argument(
        "--output-docx",
        type=Path,
        default=PROJECT_ROOT / "docs" / "manuscript_communications_biology_bilingual_with_figures_tables.docx",
    )
    parser.add_argument("--table-dir", type=Path, default=PROJECT_ROOT / "results" / "tables")
    parser.add_argument("--figure-dir", type=Path, default=PROJECT_ROOT / "results" / "figures")
    return parser.parse_args()


def clean_inline_markdown(text: str) -> str:
    text = text.strip()
    text = text.replace("**", "")
    text = text.replace("`", "")
    text = text.replace("*", "")
    text = text.replace("–", "-")
    text = re.sub(r"\s+", " ", text)
    return text


def parse_markdown_blocks(markdown_text: str) -> list[MarkdownBlock]:
    blocks: list[MarkdownBlock] = []
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            blocks.append(MarkdownBlock("h1", clean_inline_markdown(line[2:])))
        elif line.startswith("## "):
            blocks.append(MarkdownBlock("h2", clean_inline_markdown(line[3:])))
        elif line.startswith("### "):
            blocks.append(MarkdownBlock("h3", clean_inline_markdown(line[4:])))
        elif line.startswith("- "):
            blocks.append(MarkdownBlock("li", clean_inline_markdown(line[2:])))
        elif re.match(r"^\d+\.\s", line):
            blocks.append(MarkdownBlock("num", clean_inline_markdown(line)))
        else:
            blocks.append(MarkdownBlock("p", clean_inline_markdown(line)))
    return blocks


def _is_heading(block: MarkdownBlock | None) -> bool:
    return block is not None and block.kind in {"h1", "h2", "h3"}


def _is_compatible(left: MarkdownBlock | None, right: MarkdownBlock | None) -> bool:
    if left is None or right is None:
        return False
    if _is_heading(left) or _is_heading(right):
        return left.kind == right.kind
    return not _is_heading(left) and not _is_heading(right)


def pair_bilingual_blocks(
    english_blocks: list[MarkdownBlock],
    chinese_blocks: list[MarkdownBlock],
) -> list[tuple[MarkdownBlock | None, MarkdownBlock | None]]:
    pairs: list[tuple[MarkdownBlock | None, MarkdownBlock | None]] = []
    left_index = 0
    right_index = 0
    while left_index < len(english_blocks) or right_index < len(chinese_blocks):
        left = english_blocks[left_index] if left_index < len(english_blocks) else None
        right = chinese_blocks[right_index] if right_index < len(chinese_blocks) else None
        if _is_compatible(left, right):
            pairs.append((left, right))
            left_index += 1
            right_index += 1
        elif _is_heading(left) and right is not None and not _is_heading(right):
            pairs.append((None, right))
            right_index += 1
        elif _is_heading(right) and left is not None and not _is_heading(left):
            pairs.append((left, None))
            left_index += 1
        else:
            pairs.append((left, right))
            if left is not None:
                left_index += 1
            if right is not None:
                right_index += 1
    return pairs


def set_run_font(run, *, size: int = 9, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_border(cell, color: str = "D9DEE8") -> None:
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def block_style(block: MarkdownBlock | None) -> tuple[int, bool, str]:
    if block is None:
        return 9, False, "FFFFFF"
    if block.kind == "h1":
        return 12, True, "E8EEF7"
    if block.kind == "h2":
        return 11, True, "F1F4F8"
    if block.kind == "h3":
        return 10, True, "F7F9FC"
    if block.kind == "num":
        return 8, False, "FFFFFF"
    return 9, False, "FFFFFF"


def add_block_to_cell(cell, block: MarkdownBlock | None, *, side: str) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.05
    if block is None:
        run = paragraph.add_run("")
        set_run_font(run, size=9)
        return
    prefix = ""
    if block.kind == "li":
        prefix = "- "
    size, bold, fill = block_style(block)
    set_cell_shading(cell, fill)
    run = paragraph.add_run(prefix + block.text)
    color = RGBColor(55, 65, 81) if side == "en" else RGBColor(31, 41, 55)
    set_run_font(run, size=size, bold=bold, color=color)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal_style.font.size = Pt(9)


def add_title_page(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Bilingual Manuscript Review Version / 中英文对照审阅版")
    set_run_font(run, size=16, bold=True, color=RGBColor(31, 41, 55))

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Specificity-audited public spatial transcriptomics prioritization in early LUAD"
    )
    set_run_font(run, size=11, color=RGBColor(75, 85, 99))

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "正文为英文/中文逐块对照；图件和关键表格位于正文后。"
        " 术语边界：score-level in-silico target prioritization and GRN-level virtual perturbation prioritization are target-ranking layers, not causal validation."
    )
    set_run_font(run, size=9, color=RGBColor(75, 85, 99))


def add_bilingual_body(document: Document, english_text: str, chinese_text: str) -> None:
    english_blocks = parse_markdown_blocks(english_text)
    chinese_blocks = parse_markdown_blocks(chinese_text)
    pairs = pair_bilingual_blocks(english_blocks, chinese_blocks)

    document.add_heading("Manuscript Text / 正文对照", level=1)
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for cell, label in zip(header_cells, ("English", "中文"), strict=True):
        cell.text = label
        set_cell_shading(cell, "DCE6F2")
        set_cell_border(cell)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size=10, bold=True, color=RGBColor(31, 41, 55))

    for left, right in pairs:
        row = table.add_row()
        row.cells[0].width = Inches(5.25)
        row.cells[1].width = Inches(5.25)
        set_cell_border(row.cells[0])
        set_cell_border(row.cells[1])
        add_block_to_cell(row.cells[0], left, side="en")
        add_block_to_cell(row.cells[1], right, side="zh")


def add_bilingual_figures(document: Document, figure_dir: Path) -> None:
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Figures / 图件", level=1)
    for figure in BILINGUAL_FIGURES:
        path = figure_dir / figure.filename
        heading = document.add_heading(f"{figure.label_en} / {figure.label_zh}", level=2)
        for run in heading.runs:
            set_run_font(run, size=11, bold=True)
        if path.exists():
            document.add_picture(str(path), width=Inches(8.9))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph = document.add_paragraph(f"[Missing figure file: {path}]")
            for run in paragraph.runs:
                set_run_font(run, size=9, bold=True, color=RGBColor(180, 0, 0))
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(f"{figure.label_en}. {figure.caption_en} / {figure.label_zh}. {figure.caption_zh}")
        set_run_font(run, size=9, bold=True, color=RGBColor(31, 41, 55))


def add_bilingual_tables(document: Document, table_dir: Path) -> None:
    document.add_section(WD_SECTION.NEW_PAGE)
    heading = document.add_heading("Tables / 表格", level=1)
    for run in heading.runs:
        set_run_font(run, size=13, bold=True)
    note = document.add_paragraph(
        "The following tables reuse the project table-export logic so that this bilingual review file remains consistent with the Chinese Word draft. "
        "以下表格复用既有表格导出逻辑，便于与中文版 Word 稿保持一致。"
    )
    for run in note.runs:
        set_run_font(run, size=9, color=RGBColor(75, 85, 99))

    document.add_heading("Table 1 / 表 1. Public dataset composition and evidence roles / 公开数据集构成和证据角色", level=2)
    composition = zh_export.load_csv(table_dir / "figure1_dataset_composition_source.csv")
    zh_export.add_dataframe_table(
        document,
        composition,
        ["series_accession", "modality", "role", "data_extent", "n_samples", "Normal", "AAH", "AIS", "MIA", "LUAD", "IAC"],
        [
            "Dataset / 数据集",
            "Modality / 类型",
            "Evidence role / 证据角色",
            "Data extent / 数据范围",
            "Samples / 样本数",
            "Normal",
            "AAH",
            "AIS",
            "MIA",
            "LUAD",
            "IAC",
        ],
    )

    document.add_heading("Table 2 / 表 2. Integrated candidate-axis prioritization / 综合候选轴优先级排序", level=2)
    ranking = zh_export.load_csv(table_dir / "main_axis_evidence_matrix.csv").copy()
    ranking["priority_score"] = ranking["priority_score"].astype(float).map(lambda value: f"{value:.3f}")
    zh_export.add_dataframe_table(
        document,
        ranking,
        ["rank", "axis_id", "evidence_grade", "priority_score", "top_perturbed_genes", "interpretation"],
        [
            "Rank / 排名",
            "Candidate axis / 候选轴",
            "Evidence grade / 证据等级",
            "Priority score / 优先级得分",
            "Main genes / 主要基因",
            "Interpretation / 解释",
        ],
    )

    document.add_heading("Table 3 / 表 3. GSE307534 paired-patient spatial progression statistics / 配对患者空间进展统计", level=2)
    paired = zh_export.load_csv(table_dir / "gse307534_candidate_axis_paired_patient_stats.csv").copy()
    paired["paired_difference_mean"] = paired["paired_difference_mean"].astype(float).map(lambda value: f"{value:.3f}")
    paired["ci_95"] = paired.apply(
        lambda row: f"{float(row['ci_95_low']):.3f} to {float(row['ci_95_high']):.3f}",
        axis=1,
    )
    paired["positive_fraction"] = paired["positive_fraction"].astype(float).map(lambda value: f"{value:.2f}")
    paired["wilcoxon_q_bh"] = paired["wilcoxon_q_bh"].astype(float).map(lambda value: f"{value:.3g}")
    zh_export.add_dataframe_table(
        document,
        paired,
        [
            "axis_id",
            "evidence_type",
            "n_paired_patients",
            "paired_difference_mean",
            "ci_95",
            "positive_fraction",
            "wilcoxon_q_bh",
        ],
        [
            "Candidate axis / 候选轴",
            "Evidence type / 证据类型",
            "Paired patients / 配对患者数",
            "Mean delta / 平均差值",
            "95% CI",
            "Positive fraction / 正向比例",
            "BH q",
        ],
    )

    document.add_heading("Table 4 / 表 4. Score-level in-silico target prioritization / 评分层面的靶点优先级排序", level=2)
    perturb = zh_export.load_csv(table_dir / "gse307534_continuous_perturbation_mia_luad_ranking.csv").copy()
    perturb = perturb[
        perturb["perturbation_type"].eq("gene") & perturb["perturbation_factor"].astype(str).eq("0.0")
    ].copy()
    keep = ["MIF", "CD74", "CD44", "CXCR4", "SPP1", "TREM2", "PLA2G7"]
    perturb = perturb[perturb["perturbed_genes"].isin(keep)].copy()
    perturb["coupling_relative_delta_mean"] = perturb["coupling_relative_delta_mean"].astype(float).map(
        lambda value: f"{value:.3f}"
    )
    perturb["continuous_priority_score"] = perturb["continuous_priority_score"].astype(float).map(
        lambda value: f"{value:.3f}"
    )
    zh_export.add_dataframe_table(
        document,
        perturb,
        ["perturbed_genes", "axis_id", "evidence_type", "n_samples", "coupling_relative_delta_mean", "continuous_priority_score"],
        [
            "Gene / 基因",
            "Candidate axis / 候选轴",
            "Evidence type / 证据类型",
            "Samples / 样本数",
            "Relative coupling change / 相对耦合变化",
            "Priority score / 优先级得分",
        ],
    )

    document.add_heading("Table 5 / 表 5. Supplementary table manifest / 补充表清单", level=2)
    manifest = zh_export.load_csv(PROJECT_ROOT / "results" / "supplementary_tables" / "supplementary_table_manifest.csv")
    zh_export.add_dataframe_table(
        document,
        manifest,
        ["table_id", "description", "n_rows", "n_columns", "output_path"],
        ["ID", "Description / 内容", "Rows / 行数", "Columns / 列数", "File / 文件"],
    )


def main() -> int:
    args = parse_args()
    english_text = args.english_md.read_text(encoding="utf-8")
    chinese_text = args.chinese_md.read_text(encoding="utf-8")

    document = Document()
    configure_document(document)
    add_title_page(document)
    add_bilingual_body(document, english_text, chinese_text)
    add_bilingual_figures(document, args.figure_dir)
    add_bilingual_tables(document, args.table_dir)

    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output_docx)
    print(f"Wrote bilingual Word manuscript: {args.output_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
