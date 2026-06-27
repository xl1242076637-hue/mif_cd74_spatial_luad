#!/usr/bin/env python
"""Export the English Communications Biology manuscript to Word."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import NamedTuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor


SCRIPT_PATH = Path(__file__).resolve()
PROJECT_ROOT = SCRIPT_PATH.parents[1]
SCRIPTS_DIR = SCRIPT_PATH.parent
if str(SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPTS_DIR))

import export_word_manuscript as word_export  # noqa: E402


class EnglishFigure(NamedTuple):
    label: str
    caption: str
    filename: str


class EnglishTable(NamedTuple):
    label: str
    title: str


ENGLISH_FIGURES = [
    EnglishFigure(
        "Figure 1",
        "Public multi-cohort framework for early LUAD epithelial-myeloid niche mining.",
        "nature_redesign/nature_figure1_workflow_dataset_composition.png",
    ),
    EnglishFigure(
        "Figure 2",
        "Integrated evidence prioritizes MIF-CD74/CXCR4 over broader macrophage-state axes.",
        "nature_redesign/nature_figure2_axis_evidence_perturbation.png",
    ),
    EnglishFigure(
        "Figure 3",
        "Specificity auditing reframes the original SPP1 macrophage niche hypothesis.",
        "nature_redesign/nature_figure3_specificity_audit.png",
    ),
    EnglishFigure(
        "Figure 4",
        "Patient-aware spatial progression supports source-side MIF enrichment.",
        "nature_redesign/nature_figure4_spatial_axis_progression.png",
    ),
    EnglishFigure(
        "Figure 5",
        "Score-level in-silico target prioritization ranks receptor-side CD74.",
        "nature_redesign/nature_figure5_virtual_perturbation_priority.png",
    ),
    EnglishFigure(
        "Supplementary Figure 1",
        "Expression-matched and tissue-density controls for source-side MIF.",
        "supplementary_figure_mif_spatial_controls.png",
    ),
    EnglishFigure(
        "Supplementary Figure 2",
        "Focused orthogonal expression support in snRNA and bulk progression cohorts.",
        "supplementary_figure_focused_orthogonal_validation.png",
    ),
    EnglishFigure(
        "Supplementary Figure 3",
        "Specificity refinement of the original SPP1 macrophage signature.",
        "supplementary_figure_spp1_signature_refinement.png",
    ),
    EnglishFigure(
        "Supplementary Figure 4",
        "GRN-level virtual perturbation prioritization in GSE308103.",
        "supplementary_figure_grn_virtual_perturbation.png",
    ),
]


ENGLISH_TABLES = [
    EnglishTable("Table 1", "Public dataset composition and evidence roles"),
    EnglishTable("Table 2", "Integrated candidate-axis prioritization"),
    EnglishTable("Table 3", "GSE307534 paired-patient spatial progression statistics"),
    EnglishTable("Table 4", "score-level in-silico target prioritization continuous-coupling ranking"),
    EnglishTable("Table 5", "Supplementary table manifest"),
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--input-md",
        type=Path,
        default=PROJECT_ROOT / "docs" / "manuscript_communications_biology_draft.md",
    )
    parser.add_argument(
        "--output-docx",
        type=Path,
        default=PROJECT_ROOT / "docs" / "manuscript_communications_biology_draft_en_with_figures_tables.docx",
    )
    parser.add_argument("--table-dir", type=Path, default=PROJECT_ROOT / "results" / "tables")
    parser.add_argument("--figure-dir", type=Path, default=PROJECT_ROOT / "results" / "figures")
    return parser.parse_args()


def add_figures(document: Document, figure_dir: Path) -> None:
    document.add_heading("Figures", level=1)
    for figure in ENGLISH_FIGURES:
        path = figure_dir / figure.filename
        document.add_heading(figure.label, level=2)
        if path.exists():
            document.add_picture(str(path), width=Inches(6.5))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph = document.add_paragraph(f"[Missing figure file: {path}]")
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(180, 0, 0)
        paragraph = document.add_paragraph(f"{figure.label}. {figure.caption}")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            word_export.set_run_font(run, size=9, bold=True)


def add_tables(document: Document, table_dir: Path) -> None:
    document.add_heading("Tables", level=1)

    document.add_heading(f"{ENGLISH_TABLES[0].label}. {ENGLISH_TABLES[0].title}", level=2)
    composition = word_export.load_csv(table_dir / "figure1_dataset_composition_source.csv")
    word_export.add_dataframe_table(
        document,
        composition,
        ["series_accession", "modality", "role", "data_extent", "n_samples", "Normal", "AAH", "AIS", "MIA", "LUAD", "IAC"],
        ["Dataset", "Modality", "Evidence role", "Data extent", "Samples", "Normal", "AAH", "AIS", "MIA", "LUAD", "IAC"],
    )

    document.add_heading(f"{ENGLISH_TABLES[1].label}. {ENGLISH_TABLES[1].title}", level=2)
    ranking = word_export.load_csv(table_dir / "main_axis_evidence_matrix.csv").copy()
    ranking["priority_score"] = ranking["priority_score"].astype(float).map(lambda value: f"{value:.3f}")
    word_export.add_dataframe_table(
        document,
        ranking,
        ["rank", "axis_id", "evidence_grade", "priority_score", "top_perturbed_genes", "interpretation"],
        ["Rank", "Candidate axis", "Evidence grade", "Priority score", "Main genes", "Interpretation"],
    )

    document.add_heading(f"{ENGLISH_TABLES[2].label}. {ENGLISH_TABLES[2].title}", level=2)
    paired = word_export.load_csv(table_dir / "gse307534_candidate_axis_paired_patient_stats.csv").copy()
    paired["paired_difference_mean"] = paired["paired_difference_mean"].astype(float).map(lambda value: f"{value:.3f}")
    paired["ci_95"] = paired.apply(
        lambda row: f"{float(row['ci_95_low']):.3f} to {float(row['ci_95_high']):.3f}",
        axis=1,
    )
    paired["positive_fraction"] = paired["positive_fraction"].astype(float).map(lambda value: f"{value:.2f}")
    paired["wilcoxon_q_bh"] = paired["wilcoxon_q_bh"].astype(float).map(lambda value: f"{value:.3g}")
    word_export.add_dataframe_table(
        document,
        paired,
        [
            "axis_id",
            "evidence_type",
            "n_paired_patients",
            "paired_difference_mean",
            "ci_95",
            "positive_fraction",
            "wilcoxon_q_bh",
        ],
        ["Candidate axis", "Evidence type", "Paired patients", "Mean delta", "95% CI", "Positive fraction", "BH q"],
    )

    document.add_heading(f"{ENGLISH_TABLES[3].label}. {ENGLISH_TABLES[3].title}", level=2)
    perturb = word_export.load_csv(table_dir / "gse307534_continuous_perturbation_mia_luad_ranking.csv").copy()
    perturb = perturb[
        perturb["perturbation_type"].eq("gene") & perturb["perturbation_factor"].astype(str).eq("0.0")
    ].copy()
    keep = ["MIF", "CD74", "CD44", "CXCR4", "SPP1", "TREM2", "PLA2G7"]
    perturb = perturb[perturb["perturbed_genes"].isin(keep)].copy()
    perturb["coupling_relative_delta_mean"] = perturb["coupling_relative_delta_mean"].astype(float).map(
        lambda value: f"{value:.3f}"
    )
    perturb["continuous_priority_score"] = perturb["continuous_priority_score"].astype(float).map(
        lambda value: f"{value:.3f}"
    )
    word_export.add_dataframe_table(
        document,
        perturb,
        ["perturbed_genes", "axis_id", "evidence_type", "n_samples", "coupling_relative_delta_mean", "continuous_priority_score"],
        ["Gene", "Candidate axis", "Evidence type", "Samples", "Relative coupling change", "Priority score"],
    )

    document.add_heading(f"{ENGLISH_TABLES[4].label}. {ENGLISH_TABLES[4].title}", level=2)
    manifest = word_export.load_csv(PROJECT_ROOT / "results" / "supplementary_tables" / "supplementary_table_manifest.csv")
    word_export.add_dataframe_table(
        document,
        manifest,
        ["table_id", "description", "n_rows", "n_columns", "output_path"],
        ["ID", "Description", "Rows", "Columns", "File"],
    )


def main() -> int:
    args = parse_args()
    markdown_text = args.input_md.read_text(encoding="utf-8")

    document = Document()
    word_export.configure_document(document)
    word_export.add_markdown_body(document, markdown_text)
    document.add_section(WD_SECTION.NEW_PAGE)
    add_figures(document, args.figure_dir)
    document.add_section(WD_SECTION.NEW_PAGE)
    add_tables(document, args.table_dir)

    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output_docx)
    print(f"Wrote English Word manuscript: {args.output_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
