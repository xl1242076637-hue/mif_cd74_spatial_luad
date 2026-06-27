#!/usr/bin/env python
"""Export the Chinese review manuscript to Word with figures, tables and references."""

from __future__ import annotations

import argparse
import csv
import json
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SCRIPT_PATH = Path(__file__).resolve()
PROJECT_ROOT = SCRIPT_PATH.parents[1]


FIGURES = [
    (
        "图 1",
        "公开多队列框架用于早期 LUAD 上皮-髓系 niche 挖掘。",
        "nature_redesign/nature_figure1_workflow_dataset_composition.png",
    ),
    (
        "图 2",
        "综合证据优先支持 MIF-CD74/CXCR4，而不是更宽泛的巨噬细胞状态轴。",
        "nature_redesign/nature_figure2_axis_evidence_perturbation.png",
    ),
    (
        "图 3",
        "特异性审计重塑原始 SPP1 巨噬细胞 niche 假设。",
        "nature_redesign/nature_figure3_specificity_audit.png",
    ),
    (
        "图 4",
        "患者感知空间进展支持 source-side MIF 富集。",
        "nature_redesign/nature_figure4_spatial_axis_progression.png",
    ),
    (
        "图 5",
        "评分层面的 in-silico 靶点优先级排序显示受体侧 CD74 排名最高。",
        "nature_redesign/nature_figure5_virtual_perturbation_priority.png",
    ),
    (
        "补充图 1",
        "source-side MIF 的表达量匹配和组织密度对照。",
        "supplementary_figure_mif_spatial_controls.png",
    ),
    (
        "补充图 2",
        "snRNA 和 bulk 进展队列中的重点正交表达支持。",
        "supplementary_figure_focused_orthogonal_validation.png",
    ),
    (
        "补充图 3",
        "原始 SPP1 巨噬细胞 signature 的特异性修正。",
        "supplementary_figure_spp1_signature_refinement.png",
    ),
    (
        "补充图 4",
        "GSE308103 GRN-level virtual perturbation prioritization 分层候选靶点相关表达邻域。",
        "supplementary_figure_grn_virtual_perturbation.png",
    ),
]


ZOTERO_ITEM_KEYS = {
    "1": "RSB5VCKB",
    "2": "V62TJ7P9",
    "3": "U2PZUERD",
    "4": "NNXXAWI9",
    "5": "CEVHPXCD",
    "6": "SHWVSCTA",
    "7": "5P82I44J",
    "8": "T8XW7JHQ",
    "9": "H7P5UV2R",
    "10": "PJPINGB7",
    "11": "6ZUIIGFI",
    "12": "UX8MHCSG",
    "13": "DP8VWXCK",
    "14": "64C8ZTG2",
    "15": "ETPPHQ6Q",
    "16": "RAWE568I",
    "17": "5363UX9H",
    "18": "HWA9G5TF",
    "19": "RD37JPAR",
    "20": "WJ2MU6AG",
    "21": "QD9RT5DF",
    "22": "X994DRQ3",
    "23": "CIHR6KN3",
    "24": "CCTKM9R3",
    "25": "XWXQNFM6",
    "26": "IAEA39S9",
    "27": "UDAPQ38B",
    "28": "4ZPP2Z5Q",
    "29": "Z8KNSNFD",
    "30": "QVBP8TW2",
}

ZOTERO_URI_USER_ID = "13378991"


REFERENCE_FALLBACKS = {
    "10.1016/j.ccell.2025.10.004": {
        "title": "Multimodal spatial-omics reveal co-evolution of alveolar progenitors and proinflammatory niches in progression of lung precursor lesions",
        "journal": "Cancer Cell",
        "year": "2026",
        "volume": "44",
        "issue": "2",
        "pages": "321-339.e13",
        "authors": ["Peng, F.", "Sinjab, A.", "Dai, Y."],
    },
    "10.1038/s12276-022-00896-9": {
        "title": "Delineating the dynamic evolution from preneoplasia to invasive lung adenocarcinoma by integrating single-cell RNA sequencing and spatial transcriptomics",
        "journal": "Experimental & Molecular Medicine",
        "year": "2022",
        "volume": "54",
        "issue": "11",
        "pages": "2060-2076",
        "authors": ["Zhu, J.", "Fan, Y.", "Xiong, Y."],
    },
    "10.1016/j.labinv.2022.100034": {
        "title": "Tumor necrosis factor alpha-dependent lung inflammation promotes the progression of lung adenocarcinoma originating from alveolar type II cells by upregulating MIF-CD74",
        "journal": "Laboratory Investigation",
        "year": "2023",
        "volume": "103",
        "issue": "3",
        "pages": "100034",
        "authors": ["Cao, L.", "Wang, X.", "Liu, X."],
    },
    "10.1016/j.immuni.2006.08.020": {
        "title": "CD44 is the signaling component of the macrophage migration inhibitory factor-CD74 receptor complex",
        "journal": "Immunity",
        "year": "2006",
        "volume": "25",
        "issue": "4",
        "pages": "595-606",
        "authors": ["Shi, X.", "Leng, L.", "Wang, T."],
    },
    "10.1016/j.febslet.2009.07.058": {
        "title": "A functional heteromeric MIF receptor formed by CD74 and CXCR4",
        "journal": "FEBS Letters",
        "year": "2009",
        "volume": "583",
        "issue": "17",
        "pages": "2749-2757",
        "authors": ["Schwartz, V.", "Lue, H.", "Kraemer, S."],
    },
    "10.1016/j.isci.2023.107699": {
        "title": "Spatial downregulation of CD74 signatures may drive invasive component development in part-solid lung adenocarcinoma",
        "journal": "iScience",
        "year": "2023",
        "volume": "26",
        "issue": "10",
        "pages": "107699",
        "authors": ["Zhang, J.", "Zhang, J.", "Wang, S."],
    },
    "10.1007/s00262-022-03173-w": {
        "title": "Immunosuppressive TREM2(+) macrophages are associated with undesirable prognosis and responses to anti-PD-1 immunotherapy in non-small cell lung cancer",
        "journal": "Cancer Immunology, Immunotherapy",
        "year": "2022",
        "volume": "71",
        "issue": "10",
        "pages": "2511-2522",
        "authors": ["Zhang, H.", "Liu, Z.", "Wen, H."],
    },
    "10.21037/jtd-23-1012": {
        "title": "Dynamic changes in macrophage subtypes during lung cancer progression and metastasis at single-cell resolution",
        "journal": "Journal of Thoracic Disease",
        "year": "2023",
        "volume": "15",
        "issue": "8",
        "pages": "4456-4471",
        "authors": ["Wang, J.", "Wu, W.", "Xia, J."],
    },
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--input-md",
        type=Path,
        default=PROJECT_ROOT / "docs" / "manuscript_communications_biology_draft_zh.md",
    )
    parser.add_argument(
        "--output-docx",
        type=Path,
        default=PROJECT_ROOT / "docs" / "manuscript_communications_biology_draft_zh_with_figures_tables.docx",
    )
    parser.add_argument(
        "--zotero-field-docx",
        type=Path,
        default=PROJECT_ROOT / "docs" / "manuscript_communications_biology_draft_zh_zotero_fields.docx",
    )
    parser.add_argument(
        "--reference-ris",
        type=Path,
        default=PROJECT_ROOT / "results" / "references" / "communications_biology_references.ris",
    )
    parser.add_argument("--table-dir", type=Path, default=PROJECT_ROOT / "results" / "tables")
    parser.add_argument("--figure-dir", type=Path, default=PROJECT_ROOT / "results" / "figures")
    return parser.parse_args()


def set_run_font(run, *, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, *, size: int = 10) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size)


def clean_inline_markdown(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    text = text.replace("–", "-")
    return text


def parse_author(author: str) -> dict[str, str]:
    if "," in author:
        family, given = [part.strip() for part in author.split(",", 1)]
        return {"family": family, "given": given}
    parts = author.split()
    if not parts:
        return {"literal": author}
    return {"family": parts[0], "given": " ".join(parts[1:])}


def reference_metadata(row: dict[str, str]) -> dict[str, str | list[str]]:
    fallback = REFERENCE_FALLBACKS.get(row["doi"], {})
    authors = row.get("authors") or ""
    return {
        "title": row.get("title") or fallback.get("title", row["citation_short"]),
        "journal": row.get("journal") or fallback.get("journal", ""),
        "year": row.get("year") or fallback.get("year", ""),
        "volume": row.get("volume") or fallback.get("volume", ""),
        "issue": row.get("issue") or fallback.get("issue", ""),
        "pages": row.get("pages_or_article") or fallback.get("pages", ""),
        "authors": [part.strip() for part in authors.split(";") if part.strip()]
        or fallback.get("authors", []),
    }


def csl_item_data(reference_number: str, reference_rows: dict[str, dict[str, str]]) -> dict[str, object]:
    row = reference_rows[reference_number]
    metadata = reference_metadata(row)
    item: dict[str, object] = {
        "id": f"LUADREF{reference_number}",
        "type": "article-journal",
        "title": metadata["title"],
        "container-title": metadata["journal"],
        "issued": {"date-parts": [[int(str(metadata["year"]))]]},
        "DOI": row["doi"],
        "URL": row["url"],
    }
    if metadata.get("volume"):
        item["volume"] = metadata["volume"]
    if metadata.get("issue"):
        item["issue"] = metadata["issue"]
    if metadata.get("pages"):
        item["page"] = metadata["pages"]
    if metadata.get("authors"):
        item["author"] = [parse_author(author) for author in metadata["authors"]]
    return item


def expand_reference_marker(marker: str) -> list[str]:
    numbers: list[str] = []
    for part in marker.split(","):
        part = part.strip()
        if "-" in part:
            start_text, end_text = part.split("-", 1)
            numbers.extend(str(number) for number in range(int(start_text), int(end_text) + 1))
        else:
            numbers.append(part)
    return numbers


def build_zotero_citation(
    marker: str,
    citation_index: int,
    reference_rows: dict[str, dict[str, str]],
) -> str:
    citation_items = []
    for reference_number in expand_reference_marker(marker):
        item_data = csl_item_data(reference_number, reference_rows)
        citation_item: dict[str, object] = {
            "id": item_data["id"],
            "itemData": item_data,
        }
        if reference_number in ZOTERO_ITEM_KEYS:
            citation_item["uris"] = [
                f"http://zotero.org/users/{ZOTERO_URI_USER_ID}/items/{ZOTERO_ITEM_KEYS[reference_number]}"
            ]
        citation_items.append(citation_item)
    citation = {
        "citationID": f"luad-citation-{citation_index:03d}",
        "properties": {
            "formattedCitation": f"[{marker}]",
            "plainCitation": f"[{marker}]",
            "noteIndex": 0,
        },
        "citationItems": citation_items,
        "schema": "https://github.com/citation-style-language/schema/raw/master/csl-citation.json",
    }
    return " ADDIN ZOTERO_ITEM CSL_CITATION " + json.dumps(citation, ensure_ascii=False, separators=(",", ":"))


def add_zotero_field(paragraph, marker: str, instruction: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    for start in range(0, len(instruction), 900):
        instr_run = paragraph.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction[start : start + 900]
        instr_run._r.append(instr)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result_run = paragraph.add_run(f"[{marker}]")
    set_run_font(result_run, size=10)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_text_with_optional_zotero_fields(
    paragraph,
    text: str,
    *,
    reference_rows: dict[str, dict[str, str]] | None = None,
    citation_counter: list[int] | None = None,
) -> None:
    if not reference_rows or citation_counter is None:
        run = paragraph.add_run(text)
        set_run_font(run, size=10)
        return
    marker_pattern = re.compile(r"\[(\d+(?:-\d+)?(?:,\d+)*)\]")
    position = 0
    for match in marker_pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run, size=10)
        marker = match.group(1)
        citation_counter[0] += 1
        instruction = build_zotero_citation(marker, citation_counter[0], reference_rows)
        add_zotero_field(paragraph, marker, instruction)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run, size=10)


def load_reference_rows(table_dir: Path) -> dict[str, dict[str, str]]:
    rows = csv.DictReader((table_dir / "communications_biology_reference_list.csv").open(encoding="utf-8-sig"))
    return {row["reference_number"]: row for row in rows}


def set_cell_text(cell, value: object, *, bold: bool = False) -> None:
    text = "" if value is None else str(value)
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=8, bold=bold)


def add_dataframe_table(document: Document, df: pd.DataFrame, columns: list[str], headers: list[str]) -> None:
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in df[columns].itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if isinstance(value, float):
                value = f"{value:.3g}"
            set_cell_text(cells[idx], value)
    document.add_paragraph()


def load_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path, encoding="utf-8-sig")


def add_markdown_body(
    document: Document,
    markdown_text: str,
    *,
    reference_rows: dict[str, dict[str, str]] | None = None,
) -> None:
    citation_counter = [0]
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            paragraph = document.add_heading(clean_inline_markdown(line[2:]), level=0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_font(paragraph, size=16)
        elif line.startswith("## "):
            paragraph = document.add_heading(clean_inline_markdown(line[3:]), level=1)
            set_paragraph_font(paragraph, size=13)
        elif line.startswith("### "):
            paragraph = document.add_heading(clean_inline_markdown(line[4:]), level=2)
            set_paragraph_font(paragraph, size=11)
        elif line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_text_with_optional_zotero_fields(
                paragraph,
                clean_inline_markdown(line[2:]),
                reference_rows=reference_rows,
                citation_counter=citation_counter,
            )
            set_paragraph_font(paragraph)
        elif re.match(r"^\d+\.\s", line):
            paragraph = document.add_paragraph(clean_inline_markdown(line))
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            set_paragraph_font(paragraph, size=9)
        else:
            paragraph = document.add_paragraph()
            add_text_with_optional_zotero_fields(
                paragraph,
                clean_inline_markdown(line),
                reference_rows=reference_rows,
                citation_counter=citation_counter,
            )
            paragraph.paragraph_format.first_line_indent = Inches(0.28)
            paragraph.paragraph_format.line_spacing = 1.15
            set_paragraph_font(paragraph)


def add_note_box(document: Document, *, zotero_fields: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run("导出说明：")
    set_run_font(run, bold=True)
    if zotero_fields:
        message = (
            "本文档由脚本自动生成，文内引用已尝试写入 ADDIN ZOTERO_ITEM CSL_CITATION 字段代码，"
            "并嵌入 CSL itemData。若 Zotero Word 插件能识别这些字段，可在 Word 中刷新引用。"
            "若部分条目未在 Zotero 库中，请先导入随文生成的 RIS 文件。"
        )
    else:
        message = (
            "本文档由脚本自动生成，文内编号引用和参考文献表基于项目 DOI/Zotero-ready 引用清单。"
            "当前版本为可读编号引用。已另存 RIS 文件，可导入 Zotero 后用插件替换为动态引用。"
        )
    run = paragraph.add_run(message)
    set_run_font(run, size=9)
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.right_indent = Inches(0.2)


def add_figures(document: Document, figure_dir: Path) -> None:
    document.add_heading("图件", level=1)
    for figure_id, caption, filename in FIGURES:
        path = figure_dir / filename
        document.add_heading(figure_id, level=2)
        if path.exists():
            document.add_picture(str(path), width=Inches(6.5))
            last = document.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph = document.add_paragraph(f"[缺失图像文件：{path}]")
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(180, 0, 0)
        paragraph = document.add_paragraph(f"{figure_id}. {caption}")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_font(run, size=9, bold=True)


def add_tables(document: Document, table_dir: Path) -> None:
    document.add_heading("表格", level=1)

    document.add_heading("表 1. 公开数据集构成和证据角色", level=2)
    composition = load_csv(table_dir / "figure1_dataset_composition_source.csv")
    add_dataframe_table(
        document,
        composition,
        ["series_accession", "modality", "role", "data_extent", "n_samples", "Normal", "AAH", "AIS", "MIA", "LUAD", "IAC"],
        ["数据集", "类型", "证据角色", "数据范围", "样本数", "Normal", "AAH", "AIS", "MIA", "LUAD", "IAC"],
    )

    document.add_heading("表 2. 综合候选轴优先级排序", level=2)
    ranking = load_csv(table_dir / "main_axis_evidence_matrix.csv").copy()
    ranking["priority_score"] = ranking["priority_score"].astype(float).map(lambda x: f"{x:.3f}")
    add_dataframe_table(
        document,
        ranking,
        ["rank", "axis_id", "evidence_grade", "priority_score", "top_perturbed_genes", "interpretation"],
        ["排名", "候选轴", "证据等级", "优先级得分", "主要基因", "解释"],
    )

    document.add_heading("表 3. GSE307534 配对患者空间进展统计", level=2)
    paired = load_csv(table_dir / "gse307534_candidate_axis_paired_patient_stats.csv").copy()
    paired["paired_difference_mean"] = paired["paired_difference_mean"].astype(float).map(lambda x: f"{x:.3f}")
    paired["ci_95"] = paired.apply(
        lambda row: f"{float(row['ci_95_low']):.3f} to {float(row['ci_95_high']):.3f}",
        axis=1,
    )
    paired["positive_fraction"] = paired["positive_fraction"].astype(float).map(lambda x: f"{x:.2f}")
    paired["wilcoxon_q_bh"] = paired["wilcoxon_q_bh"].astype(float).map(lambda x: f"{x:.3g}")
    add_dataframe_table(
        document,
        paired,
        [
            "axis_id",
            "evidence_type",
            "n_paired_patients",
            "paired_difference_mean",
            "ci_95",
            "positive_fraction",
            "wilcoxon_q_bh",
        ],
        ["候选轴", "证据类型", "配对患者数", "平均差值", "95% CI", "正向比例", "BH q"],
    )

    document.add_heading("表 4. 评分层面的 in-silico 靶点优先级排序（连续耦合）", level=2)
    perturb = load_csv(table_dir / "gse307534_continuous_perturbation_mia_luad_ranking.csv").copy()
    perturb = perturb[
        perturb["perturbation_type"].eq("gene") & perturb["perturbation_factor"].astype(str).eq("0.0")
    ].copy()
    keep = ["MIF", "CD74", "CD44", "CXCR4", "SPP1", "TREM2", "PLA2G7"]
    perturb = perturb[perturb["perturbed_genes"].isin(keep)].copy()
    perturb["coupling_relative_delta_mean"] = perturb["coupling_relative_delta_mean"].astype(float).map(
        lambda x: f"{x:.3f}"
    )
    perturb["continuous_priority_score"] = perturb["continuous_priority_score"].astype(float).map(lambda x: f"{x:.3f}")
    add_dataframe_table(
        document,
        perturb,
        ["perturbed_genes", "axis_id", "evidence_type", "n_samples", "coupling_relative_delta_mean", "continuous_priority_score"],
        ["基因", "候选轴", "证据类型", "样本数", "相对耦合变化", "优先级得分"],
    )

    document.add_heading("表 5. 补充表清单", level=2)
    manifest = load_csv(PROJECT_ROOT / "results" / "supplementary_tables" / "supplementary_table_manifest.csv")
    add_dataframe_table(
        document,
        manifest,
        ["table_id", "description", "n_rows", "n_columns", "output_path"],
        ["ID", "内容", "行数", "列数", "文件"],
    )


def write_ris(reference_csv: Path, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    rows = list(csv.DictReader(reference_csv.open(encoding="utf-8-sig")))
    lines: list[str] = []
    for row in rows:
        doi = row["doi"]
        metadata = reference_metadata(row)
        lines.append("TY  - JOUR")
        for author in metadata.get("authors", []):
            lines.append(f"AU  - {author}")
        lines.append(f"TI  - {metadata['title']}")
        lines.append(f"T2  - {metadata['journal']}")
        lines.append(f"PY  - {metadata['year']}")
        if metadata.get("volume"):
            lines.append(f"VL  - {metadata['volume']}")
        if metadata.get("issue"):
            lines.append(f"IS  - {metadata['issue']}")
        if metadata.get("pages"):
            lines.append(f"SP  - {metadata['pages']}")
        lines.append(f"DO  - {doi}")
        lines.append(f"UR  - {row['url']}")
        lines.append("ER  - ")
        lines.append("")
    output_path.write_text("\n".join(lines), encoding="utf-8")


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal_style.font.size = Pt(10)


def main() -> int:
    args = parse_args()
    markdown_text = args.input_md.read_text(encoding="utf-8")

    for output_path, reference_rows, zotero_fields in (
        (args.output_docx, None, False),
        (args.zotero_field_docx, load_reference_rows(args.table_dir), True),
    ):
        document = Document()
        configure_document(document)
        add_markdown_body(document, markdown_text, reference_rows=reference_rows)
        add_note_box(document, zotero_fields=zotero_fields)
        document.add_section(WD_SECTION.NEW_PAGE)
        add_figures(document, args.figure_dir)
        document.add_section(WD_SECTION.NEW_PAGE)
        add_tables(document, args.table_dir)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

    write_ris(args.table_dir / "communications_biology_reference_list.csv", args.reference_ris)
    print(f"Wrote Word manuscript: {args.output_docx}")
    print(f"Wrote Zotero-field Word manuscript: {args.zotero_field_docx}")
    print(f"Wrote Zotero-ready RIS: {args.reference_ris}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
